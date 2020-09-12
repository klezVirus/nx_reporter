from formatter import nx_oracle
from plotter import nx_plot
from filterer import nx_filter
from configurer import nx_config
from splitter import nx_split
from docx import Document
from docx.shared import Inches
from docx.shared import Cm
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
import time

class nx_report(object):

	def __init__(self):
		self.nx_config = nx_config()
		self.nx_config.load_config()
		self.nx_config.rebase()
		self.nx_split = nx_split(self.nx_config)
		self.mode = self.nx_config.get_mode()
		return 

	def stringpath(self,ptype,pstr):
		if ptype == "csv":
			return self.nx_config.get("GLOBAL_DIRS","CSV_DIR") + "/" + pstr + "/" + pstr + ".csv"
		elif ptype == "png":
			return self.nx_config.get("GLOBAL_DIRS","PNG_DIR") + "/" + pstr + "/"
		elif ptype == "out":
			return self.nx_config.get("GLOBAL_DIRS","OUTPUT_DIR") + "/" + pstr + "/"

	def get_all_styles(self):
		styles = self.document.styles

		paragraph_styles = [ s for s in styles]# if s.type == WD_STYLE_TYPE.PARAGRAPH]

		for s in paragraph_styles:
			print (s.name)

	def write_table(self,dataframe, title="",usetitle=True):
		dataframe = dataframe.round(2)
		self.document.add_heading(title, level=2)
		gap = 0
		if usetitle:
			gap += 1

		t = self.document.add_table(dataframe.shape[0]+1, dataframe.shape[1])
		t.style = self.nx_config.get("TEMPLATE","TABLE_STYLE")

		if usetitle:
			hdr_cells = t.rows[0].cells
			cell = hdr_cells[0]
			for i in range(gap, dataframe.shape[1]):
				cell = cell.merge(hdr_cells[i])
			cell.add_paragraph(title).alignment = WD_ALIGN_PARAGRAPH.CENTER

		for j in range(dataframe.shape[-1]):
			t.cell(gap,j).text = dataframe.columns[j]

		for i in range(gap, dataframe.shape[0]):
			if title in ["Most Important Vulnerabilities", "Most Common Vulnerabilities", "Highest Risk Vulnerabilities","Top 25 Remediations by Risk"]:
				t.cell(i+1,0).width = 11* 914400 
			for j in range(dataframe.shape[-1]):
				if not j==0:
					t.cell(i+1,j).width = 914400
				t.cell(i+1,j).text = str(dataframe.values[i,j])
	
	def write_vdtable(self, dataframe):
		dataframe = dataframe.round(2)


		df1 = dataframe[["Severity","CVSS","Vector","Description","Reference"]]
		df2 = dataframe[["ID","Title", "Evidence","Ports","Assets"]]
		df3 = dataframe[["Solution"]]

		rows,columns = dataframe.shape

		for i in range( rows):
			self.document.add_heading(str(df2.values[i,1]), level=1)
			t = self.document.add_table(5, 2)
			t.style = "Medium Grid 1"

			for k in range(len(df1.columns)):
				t.cell(k,0).text = df1.columns[k]
				if df1.columns[k] == "Reference":
					t.cell(k,1).text = str(df1.values[i,k]).replace("@@", "\n")
				else:
					t.cell(k,1).text = str(df1.values[i,k])

			self.document.add_paragraph()

			t = self.document.add_table(2, df2.shape[1])
			t.style = "Light Grid"
				
			for k in range(len(df2.columns)):
				t.cell(0,k).text = str(df2.columns[k])
				if df2.columns[k] =="Evidence":
					t.cell(1,k).text = str(df2.values[i,k]).replace("*","\n") 
				elif df2.columns[k] =="Ports":
					temp = set(str(df2.values[i,k]).split("-"))
					t.cell(1,k).text = "\n".join(sorted(list(temp)))
				elif df2.columns[k] =="Assets":
					temp = set(str(df2.values[i,k]).split(";"))
					t.cell(1,k).text = "\n".join(sorted(list(temp)))
				else:
					t.cell(1,k).text = str(df2.values[i,k]) 
			
			self.document.add_paragraph()
				
			p = self.document.add_heading("Solution",level=2)
			t = self.document.add_table(1, df3.shape[1])
			for k in range(df3.shape[1]):
				t.cell(0,k).text = str(df3.values[i,k])
			self.document.add_paragraph()
				

	def write_image(self, imgpath, title):
		if title in ["Most Important Vulnerabilities", "Top 25 Remediations by Risk"]:
			return
		self.document.add_paragraph("")
		self.document.add_picture(imgpath + title + ".png", width=Inches(7.00))

	def write_report(self, imgpath, outpath, dtitle="dashboard.docx"):
		self.document = Document(self.nx_config.get("FILES","TEMPLATE"))
		if dtitle=="dashboard.docx":
			self.document.add_heading('VULNERABILITY REPORT', 0)
			self.document.add_heading('RESULT OF ACTIVITY', level=1)
			self.document.add_heading('SCAN INFORMATION', level=2)

			table = self.document.add_table(rows=6, cols=2)
			table.style = 'Medium Grid 1'
			hdr_cells = table.rows[0].cells
			cell = hdr_cells[0].merge(hdr_cells[1])
			cell.add_paragraph('Scan Information').alignment = WD_ALIGN_PARAGRAPH.CENTER
			static = ["Company", "Site", "Site Description", "Data Inizio", "Data Fine"] #, "Policy Name", "Policy Descriton", "Policy Checks"]
			i = 1
			for k in static:
				table.rows[i].cells[0].text = k
				i+=1
			
			table.rows[1].cells[1].text = "TEST COMPANY"
			table.rows[2].cells[1].text = self.nx_oracle.scaninfo["Site Name"]
			table.rows[3].cells[1].text = self.nx_oracle.scaninfo["Site Description"]
			table.rows[4].cells[1].text = self.nx_oracle.scaninfo["Start"]
			table.rows[5].cells[1].text = self.nx_oracle.scaninfo["End"]

			self.document.add_heading('RESULT INFORMATION', level=2)

			# table.rows[8].cells[1].paragraphs[0].text = "Asset Discovery"
			# table.rows[8].cells[1].paragraphs[0].style = 'List Bullet'
			# table.rows[8].cells[1].add_paragraph("Vulnerabilities", style = 'List Bullet')
			# table.rows[8].cells[1].add_paragraph("Web Spidering", style = 'List Bullet')
			# table.rows[8].cells[1].add_paragraph("Policies", style = 'List Bullet')
			for key, df in self.nx_oracle.get_tables().items():
				if not key == "Vulnerability List":
					print (df)
					self.write_image(imgpath, key)
					self.write_table(df, title=key,usetitle=False)
					
		elif dtitle == "Vulnerability Details.docx":
			self.document.add_heading('VULNERABILITY REPORT', 0)
			self.document.add_heading('RESULT OF ACTIVITY', level=1)
			self.document.add_heading('SCAN INFORMATION', level=2)

			table = self.document.add_table(rows=6, cols=2)
			table.style = 'Medium Grid 1'
			hdr_cells = table.rows[0].cells
			cell = hdr_cells[0].merge(hdr_cells[1])
			cell.add_paragraph('Scan Information').alignment = WD_ALIGN_PARAGRAPH.CENTER
			static = ["Company", "Site", "Site Description", "Data Inizio", "Data Fine"] #, "Policy Name", "Policy Descriton", "Policy Checks"]
			i = 1
			for k in static:
				table.rows[i].cells[0].text = k
				i+=1
			
			table.rows[1].cells[1].text = "TEST COMPANY"
			table.rows[2].cells[1].text = self.nx_oracle.scaninfo["Site Name"]
			table.rows[3].cells[1].text = self.nx_oracle.scaninfo["Site Description"]
			table.rows[4].cells[1].text = self.nx_oracle.scaninfo["Start"]
			table.rows[5].cells[1].text = self.nx_oracle.scaninfo["End"]

			self.document.add_heading('RESULT INFORMATION', level=2)

			# table.rows[8].cells[1].paragraphs[0].text = "Asset Discovery"
			# table.rows[8].cells[1].paragraphs[0].style = 'List Bullet'
			# table.rows[8].cells[1].add_paragraph("Vulnerabilities", style = 'List Bullet')
			# table.rows[8].cells[1].add_paragraph("Web Spidering", style = 'List Bullet')
			# table.rows[8].cells[1].add_paragraph("Policies", style = 'List Bullet')
			for key, df in self.nx_oracle.get_tables().items():
				if key in ["Vulnerability List","Asset By Risk", "Asset By Vulnerabilities", "Top 25 Remediation by Risk"]:
					continue
				self.write_image(imgpath, key)
				self.write_table(df, title=key,usetitle=False)
			
			self.write_vdtable(self.nx_oracle.get_vdtable())

		
		for section in self.document.sections:
			section.left_margin = Cm(2.0)

			section.right_margin = Cm(2.0)

		#document.add_picture('monty-truth.png', width=Inches(1.25))		

		self.document.add_page_break()

		self.document.save(outpath+dtitle)

	def gentab(self):
		self.nx_split.splitfile().writeondisk()
		file_to_parse = self.nx_split.get_splitted_files()
		for file in file_to_parse:
			print(self.stringpath("csv", file))
			self.nx_oracle = nx_oracle(self.nx_config,self.stringpath("csv",file))
			
			if self.nx_oracle.data is None:
				continue
			self.nx_filter = nx_filter()
			self.nx_oracle.generateTableData()
			self.nx_plot = nx_plot(self.nx_config,self.stringpath("png",file))
			self.nx_plot.plot_tables(self.nx_oracle.get_tables())
			self.write_report(self.stringpath("png",file),self.stringpath("out",file))
			self.write_report(self.stringpath("png",file),self.stringpath("out",file), dtitle="Vulnerability Details.docx")


reporter = nx_report()

if reporter.mode == "perpetual":
	while not reporter.nx_split.filesearch():
		time.sleep(60)
	print("FOUND")
	time.sleep(int(reporter.nx_split.wait))
elif reporter.mode == "onetime":
	reporter.gentab()